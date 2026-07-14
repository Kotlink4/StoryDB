from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

import pypdfium2 as pdfium
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPO_ROOT = Path(r"C:\Users\vovaz\source\repos\StoryDB")
DESKTOP = Path(r"C:\Users\vovaz\Desktop")
OUTPUT = REPO_ROOT / "docs" / "vkr" / "articles_results_appendix.docx"
ASSET_DIR = REPO_ROOT / "docs" / "vkr" / "articles_appendix_assets"


@dataclass(frozen=True)
class ArticleSource:
    number: str
    docx_path: Path
    pdf_path: Path
    result_summary: str
    preview_caption: str


ARTICLES = (
    ArticleSource(
        "1",
        DESKTOP / "storydb_article_1_engineering_technologies_final_short_abstract.docx",
        DESKTOP / "storydb_article_1_engineering_technologies_final.pdf",
        (
            "Подготовлена статья о веб-системе StoryDB, в которой описаны назначение "
            "системы, архитектура, модель данных и основные модули проекта."
        ),
        "Первая страница статьи о веб-системе StoryDB",
    ),
    ArticleSource(
        "2",
        DESKTOP / "storydb_article_2_engineering_technologies_antiplagiat_s_avtorami_short_abstract.docx",
        DESKTOP / "storydb_article_2_engineering_technologies_antiplagiat_s_avtorami.pdf",
        (
            "Подготовлена статья о модели связей и хронологии, в которой показано "
            "совместное использование графа, структур и таймлайна StoryDB."
        ),
        "Первая страница статьи о связях и хронологии StoryDB",
    ),
)


def set_run_font(run, name: str = "Times New Roman", size: float = 14, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_style_font(style, name: str, size: float, bold: bool = False, color: str = "000000") -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style.element.rPr.rFonts.set(qn("w:ascii"), name)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = document.styles
    normal = styles["Normal"]
    set_style_font(normal, "Times New Roman", 14)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        set_style_font(style, "Times New Roman", 14, bold=True)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.first_line_indent = Cm(1.25)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(6)

    if "CaptionText" not in [style.name for style in styles]:
        caption = styles.add_style("CaptionText", 1)
    else:
        caption = styles["CaptionText"]
    set_style_font(caption, "Times New Roman", 12)
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.line_spacing = 1
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)


def add_page_number_footer(document: Document) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr)
    run._r.append(fld_char_2)
    set_run_font(run, size=12)


def clean(text: str) -> str:
    return " ".join((text or "").split())


def display_title(title: str) -> str:
    normalized = clean(title)
    if not normalized:
        return normalized
    lowered = normalized[:1] + normalized[1:].lower()
    return lowered.replace("storydb", "StoryDB")


def display_authors(authors: str) -> str:
    return re.sub(r"(?<=\D)1\b", "", clean(authors)).replace("  ", " ")


def display_organization(organization: str) -> str:
    return re.sub(r"^\d+\s*", "", clean(organization))


def read_article(source: ArticleSource) -> dict:
    doc = Document(source.docx_path)
    paragraphs = [clean(paragraph.text) for paragraph in doc.paragraphs]
    paragraphs = [text for text in paragraphs if text]

    def find(prefix: str) -> str:
        for text in paragraphs:
            if text.startswith(prefix):
                return text
        return ""

    title_index = next(index for index, text in enumerate(paragraphs) if text.isupper() and "STORYDB" in text)
    title = paragraphs[title_index]
    authors = display_authors(paragraphs[title_index + 1])
    organization = display_organization(paragraphs[title_index + 2])
    annotation = find("Аннотация:")
    keywords = find("Ключевые слова:")

    conclusion = []
    in_conclusion = False
    for text in paragraphs:
        if text == "Заключение":
            in_conclusion = True
            continue
        if in_conclusion and text in {"Литература", "References"}:
            break
        if in_conclusion:
            conclusion.append(text)

    table_data = []
    if doc.tables:
        for row in doc.tables[0].rows:
            table_data.append([clean(cell.text) for cell in row.cells])

    return {
        "udc": paragraphs[0],
        "title": title,
        "display_title": display_title(title),
        "authors": authors,
        "organization": organization,
        "annotation": annotation,
        "keywords": keywords,
        "conclusion": conclusion,
        "table": table_data,
    }


def render_pdf_first_page(source: ArticleSource) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    out = ASSET_DIR / f"article_{source.number}_first_page.png"
    pdf = pdfium.PdfDocument(str(source.pdf_path))
    page = pdf[0]
    bitmap = page.render(scale=2.0)
    image = bitmap.to_pil()
    image.save(out)
    return out


def add_plain_paragraph(document: Document, text: str, *, bold: bool = False, align=None) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0 if align == WD_ALIGN_PARAGRAPH.CENTER else 1.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=14, bold=bold)


def add_labeled_paragraph(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=14, bold=True)
    value_run = paragraph.add_run(value)
    set_run_font(value_run, size=14)


def add_section_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    set_run_font(run, size=14, bold=True)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def format_cell(cell, *, bold: bool = False, size: float = 11, fill: str | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        set_cell_shading(cell, fill)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold)


def add_table_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="CaptionText")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run(text)


def add_figure_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="CaptionText")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(text)


def add_metadata_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [2500, 6860]
    for row_index, (label, value) in enumerate(rows):
        row = table.rows[row_index]
        row.cells[0].text = label
        row.cells[1].text = value
        for col_index, width in enumerate(widths):
            set_cell_width(row.cells[col_index], width)
            format_cell(row.cells[col_index], bold=col_index == 0, size=11, fill="F2F4F7" if col_index == 0 else None)


def add_source_table(document: Document, table_data: list[list[str]], caption: str) -> None:
    add_table_caption(document, caption)
    rows = len(table_data)
    cols = len(table_data[0])
    table = document.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    if cols == 2:
        widths = [2600, 6760]
    elif cols == 3:
        widths = [2100, 3500, 3760]
    else:
        widths = [int(9360 / cols)] * cols

    for row_index, row_values in enumerate(table_data):
        for col_index, value in enumerate(row_values):
            cell = table.cell(row_index, col_index)
            cell.text = value
            set_cell_width(cell, widths[col_index])
            format_cell(cell, bold=row_index == 0, size=10.5, fill="F2F4F7" if row_index == 0 else None)


def add_summary_table(document: Document, articles: list[tuple[ArticleSource, dict]]) -> None:
    add_table_caption(document, "Таблица Б.1 - Сведения о подготовленных статьях")
    rows = [["N", "Название статьи", "Краткий результат"]]
    for source, data in articles:
        rows.append([source.number, data["display_title"], source.result_summary])

    table = document.add_table(rows=len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [650, 3900, 4810]
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            cell.text = value
            set_cell_width(cell, widths[col_index])
            format_cell(cell, bold=row_index == 0, size=10.5, fill="F2F4F7" if row_index == 0 else None)
            if col_index == 0:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_article_block(document: Document, source: ArticleSource, data: dict, preview_path: Path) -> None:
    add_section_title(document, f"Б.{source.number} Результаты статьи {source.number}")
    add_labeled_paragraph(document, "Название", data["display_title"])
    add_labeled_paragraph(document, "Авторы", data["authors"])
    add_labeled_paragraph(document, "УДК", data["udc"])
    add_labeled_paragraph(document, "Аннотация", data["annotation"])
    add_labeled_paragraph(document, "Ключевые слова", data["keywords"])

    add_section_title(document, "Основной результат")
    add_plain_paragraph(document, source.result_summary)

    document.add_page_break()
    add_section_title(document, f"Б.{source.number}.1 Фрагмент оформленной статьи")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    run.add_picture(str(preview_path), width=Cm(14.2))
    add_figure_caption(document, f"Рисунок Б.{source.number} - {source.preview_caption}")


def add_articles_list(document: Document, articles: list[tuple[ArticleSource, dict]]) -> None:
    add_section_title(document, "Перечень подготовленных статей")
    for source, data in articles:
        add_plain_paragraph(document, f"{source.number}. {data['display_title']}. {source.result_summary}")


def build_document() -> None:
    for source in ARTICLES:
        if not source.docx_path.exists():
            raise FileNotFoundError(source.docx_path)
        if not source.pdf_path.exists():
            raise FileNotFoundError(source.pdf_path)

    document = Document()
    configure_document(document)
    add_page_number_footer(document)

    loaded = [(source, read_article(source)) for source in ARTICLES]
    previews = {source.number: render_pdf_first_page(source) for source in ARTICLES}

    add_plain_paragraph(document, "Приложение Б", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_plain_paragraph(
        document,
        "Результаты подготовки научных статей по теме выпускной квалификационной работы",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    document.add_paragraph()
    add_plain_paragraph(
        document,
        (
            "В приложении приведены сведения о двух научных статьях, подготовленных по материалам "
            "проекта StoryDB. Статьи отражают разные стороны выпускной квалификационной работы: "
            "разработку веб-системы для ведения базы знаний сюжетного проекта и моделирование "
            "связей с хронологией внутри этой системы."
        ),
    )
    add_articles_list(document, loaded)

    for index, (source, data) in enumerate(loaded):
        document.add_page_break()
        add_article_block(document, source, data, previews[source.number])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
