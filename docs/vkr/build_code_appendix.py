from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPO_ROOT = Path(r"C:\Users\vovaz\source\repos\StoryDB")
OUTPUT = REPO_ROOT / "docs" / "vkr" / "code_appendix_storydb.docx"


@dataclass(frozen=True)
class SourceRange:
    path: str
    start: int | None = None
    end: int | None = None


@dataclass(frozen=True)
class Listing:
    number: str
    title: str
    note: str
    sources: tuple[SourceRange, ...]


LISTINGS = (
    Listing(
        "А.1",
        "Сущность объекта художественного мира",
        "Фрагмент показывает базовую модель объекта, его атрибуты и связи с другими объектами.",
        (
            SourceRange(r"StoryDB.Api\Data\Entities\StoryObject.cs"),
            SourceRange(r"StoryDB.Api\Data\Entities\ObjectAttribute.cs"),
            SourceRange(r"StoryDB.Api\Data\Entities\ObjectRelation.cs"),
        ),
    ),
    Listing(
        "А.2",
        "Настраиваемые каталоги и поля",
        "Листинг отражает справочники художественного мира, записи каталогов и пользовательские значения.",
        (
            SourceRange(r"StoryDB.Api\Data\Entities\Catalog.cs"),
            SourceRange(r"StoryDB.Api\Data\Entities\CatalogEntry.cs"),
            SourceRange(r"StoryDB.Api\Data\Entities\CatalogFieldDefinition.cs"),
            SourceRange(r"StoryDB.Api\Data\Entities\CatalogEntryFieldValue.cs"),
        ),
    ),
    Listing(
        "А.3",
        "Структуры проекта",
        "Фрагмент описывает структуры, узлы и назначения объектов на элементы схемы.",
        (
            SourceRange(r"StoryDB.Api\Data\Entities\Structure.cs"),
            SourceRange(r"StoryDB.Api\Data\Entities\StructureNode.cs"),
            SourceRange(r"StoryDB.Api\Data\Entities\StructureAssignment.cs"),
        ),
    ),
    Listing(
        "А.4",
        "События и временные изменения",
        "Листинг показывает сущности таймлайна, участников событий и изменения состояния объектов.",
        (
            SourceRange(r"StoryDB.Api\Data\Entities\TimelineEvent.cs"),
            SourceRange(r"StoryDB.Api\Data\Entities\TimelineChange.cs"),
            SourceRange(r"StoryDB.Api\Data\Entities\TimelineEventLink.cs"),
        ),
    ),
    Listing(
        "А.5",
        "Контекст данных приложения",
        "В фрагменте приведена регистрация основных таблиц EF Core.",
        (SourceRange(r"StoryDB.Api\Data\StoryDbContext.cs", 1, 75),),
    ),
    Listing(
        "А.6",
        "Конфигурация таблиц объектов",
        "Фрагмент содержит связи, индексы и ограничения для объектов, атрибутов и отношений.",
        (SourceRange(r"StoryDB.Api\Data\StoryDbContext.Objects.cs", 1, 110),),
    ),
    Listing(
        "А.7",
        "Конфигурация каталогов",
        "Фрагмент показывает настройку справочников, записей, полей и значений каталогов.",
        (SourceRange(r"StoryDB.Api\Data\StoryDbContext.Catalogs.cs", 1, 105),),
    ),
    Listing(
        "А.8",
        "Конфигурация временной линии",
        "Фрагмент отражает настройку событий, связей событий, участников и временных изменений.",
        (SourceRange(r"StoryDB.Api\Data\StoryDbContext.Timelines.cs", 1, 115),),
    ),
    Listing(
        "А.9",
        "DTO и request-модели объектов",
        "Листинг показывает контракт обмена данными между клиентом и сервером для карточек объектов.",
        (SourceRange(r"StoryDB.Api\Contracts\Objects\ObjectContracts.cs", 1, 125),),
    ),
    Listing(
        "А.10",
        "Контроллер объектов",
        "Фрагмент демонстрирует REST-интерфейс для чтения, создания и обновления объектов проекта.",
        (SourceRange(r"StoryDB.Api\Controllers\ObjectsController.cs", 1, 130),),
    ),
    Listing(
        "А.11",
        "Создание и обновление объекта",
        "В листинге приведена часть сервисного слоя, отвечающая за сохранение карточки объекта.",
        (
            SourceRange(r"StoryDB.Api\Services\Objects\ObjectService.cs", 1, 105),
            SourceRange(r"StoryDB.Api\Services\Objects\ObjectService.cs", 139, 205),
        ),
    ),
    Listing(
        "А.12",
        "Синхронизация атрибутов и связей объекта",
        "Фрагмент показывает обновление связанных данных: атрибутов, каталогов, иерархий и отношений.",
        (
            SourceRange(r"StoryDB.Api\Services\Objects\ObjectService.Sync.cs", 1, 110),
            SourceRange(r"StoryDB.Api\Services\Objects\ObjectService.RelationsSync.cs", 1, 80),
        ),
    ),
    Listing(
        "А.13",
        "Сервис каталогов",
        "Листинг показывает создание и обновление записей каталога с пользовательскими полями.",
        (SourceRange(r"StoryDB.Api\Services\Catalogs\CatalogService.Entries.cs", 1, 135),),
    ),
    Listing(
        "А.14",
        "События таймлайна и раскладка",
        "Фрагмент отражает серверную логику событий и подготовку визуального размещения таймлайна.",
        (
            SourceRange(r"StoryDB.Api\Services\Timelines\TimelineService.Events.cs", 1, 120),
            SourceRange(r"StoryDB.Api\Services\Timelines\TimelineService.LayoutEngine.cs", 1, 85),
        ),
    ),
    Listing(
        "А.15",
        "Экспорт DOCX и клиентские операции",
        "Листинг объединяет фрагменты сервера экспорта и клиентского интерфейса подготовки досье.",
        (
            SourceRange(r"StoryDB.Api\Services\Exports\ProjectExportService.cs", 1, 100),
            SourceRange(r"StoryDB.Api\Services\Exports\ProjectExportJobService.cs", 33, 95),
            SourceRange(r"storydb.client\src\components\ProjectExportWorkspace.tsx", 145, 190),
            SourceRange(r"storydb.client\src\api\exports.ts", 45, 100),
        ),
    ),
)


def set_font(run, name: str, size: float, bold: bool = False, color: str | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_style_font(style, name: str, size: float, color: str | None = None, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    style.element.rPr.rFonts.set(qn("w:ascii"), name)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = document.styles
    set_style_font(styles["Normal"], "Times New Roman", 12)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    styles["Normal"].paragraph_format.space_after = Pt(6)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
        style = styles[style_name]
        set_style_font(style, "Times New Roman", size, "000000", bold=True)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(6)

    if "CodeLine" not in [style.name for style in styles]:
        code_style = styles.add_style("CodeLine", 1)
    else:
        code_style = styles["CodeLine"]
    set_style_font(code_style, "Courier New", 7.5)
    code_style.paragraph_format.space_before = Pt(0)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing = Pt(8.7)

    if "SourceLabel" not in [style.name for style in styles]:
        source_style = styles.add_style("SourceLabel", 1)
    else:
        source_style = styles["SourceLabel"]
    set_style_font(source_style, "Times New Roman", 10, "444444")
    source_style.paragraph_format.space_before = Pt(4)
    source_style.paragraph_format.space_after = Pt(2)

    if "ListingCaption" not in [style.name for style in styles]:
        caption_style = styles.add_style("ListingCaption", 1)
    else:
        caption_style = styles["ListingCaption"]
    set_style_font(caption_style, "Times New Roman", 12, "000000", bold=True)
    caption_style.paragraph_format.space_before = Pt(0)
    caption_style.paragraph_format.space_after = Pt(4)


def add_footer(document: Document) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Приложение. Фрагменты программного кода StoryDB")
    set_font(run, "Times New Roman", 9, color="555555")


def add_title(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run("Приложение А\nФрагменты программного кода веб-приложения StoryDB")
    set_font(run, "Times New Roman", 16, bold=True)

    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(
        "В приложении приведены сокращенные фрагменты программного кода, "
        "которые отражают основные части разработанного веб-приложения: "
        "информационную модель, конфигурацию базы данных, серверные сервисы, "
        "API, работу с таймлайном, экспортом DOCX и клиентским интерфейсом."
    )


def read_source(source: SourceRange) -> list[tuple[int, str]]:
    path = REPO_ROOT / source.path
    lines = path.read_text(encoding="utf-8-sig").splitlines()
    start = source.start or 1
    end = source.end or len(lines)
    return [(index, lines[index - 1]) for index in range(start, min(end, len(lines)) + 1)]


def prepare_code_line(code: str) -> str:
    return code.replace("\t", "    ").rstrip()


def add_source_label(document: Document, source: SourceRange) -> None:
    start = source.start or 1
    end = source.end or "конец файла"
    paragraph = document.add_paragraph(style="SourceLabel")
    paragraph.add_run(f"Фрагмент файла: {source.path}; строки {start}-{end}")


def add_code_line(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="CodeLine")
    run = paragraph.add_run(text)
    set_font(run, "Courier New", 7.5)


def add_listing(document: Document, listing: Listing, first: bool = False) -> None:
    if not first:
        document.add_page_break()

    caption = document.add_paragraph(style="ListingCaption")
    caption.add_run(f"Листинг {listing.number} - {listing.title}")

    note = document.add_paragraph()
    note.paragraph_format.space_after = Pt(4)
    note.add_run(listing.note)

    for source_index, source in enumerate(listing.sources):
        if source_index > 0:
            add_code_line(document, "")
        add_source_label(document, source)
        for _, code in read_source(source):
            add_code_line(document, prepare_code_line(code))


def audit_docx(path: Path) -> None:
    import zipfile

    with zipfile.ZipFile(path, "r") as archive:
        names = set(archive.namelist())
    required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
    missing = sorted(required - names)
    if missing:
        raise RuntimeError(f"Missing DOCX parts: {missing}")


def main() -> None:
    document = Document()
    configure_document(document)
    add_footer(document)
    add_title(document)
    for index, listing in enumerate(LISTINGS):
        add_listing(document, listing, first=index == 0)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    audit_docx(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
