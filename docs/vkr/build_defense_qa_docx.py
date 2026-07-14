from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("Шпаргалка_вопросы_ответы_StoryDB.docx")


SECTIONS = [
    (
        "По теме",
        [
            (
                "В чем цель работы?",
                "Разработать веб-приложение для авторов художественных произведений, позволяющее вести базу знаний проекта, хранить объекты мира, связи между ними и события временной линии.",
            ),
            (
                "В чем актуальность?",
                "При работе над крупным художественным миром материалы быстро становятся разрозненными. Обычные документы и таблицы плохо показывают связи между персонажами, местами, предметами и событиями.",
            ),
            (
                "В чем новизна?",
                "Новизна состоит в том, что художественный мир представлен как связанная и изменяющаяся во времени система: объекты, каталоги, структуры, граф связей и таймлайн объединены в одной модели.",
            ),
            (
                "Чем отличается от аналогов?",
                "Аналоги часто ориентированы на заметки, публикацию или готовые шаблоны. StoryDB делает упор на связи между сущностями, временные изменения и гибкие пользовательские каталоги.",
            ),
        ],
    ),
    (
        "По базе данных",
        [
            (
                "Как устроена БД?",
                "Центральная сущность - проект. К нему привязаны объекты, каталоги, структуры, таймлайн, медиафайлы, раскладки графов и снимки проекта.",
            ),
            (
                "Что хранится в объектах?",
                "Объекты художественного мира: персонажи, предметы, места, организации. У них есть тип, название, описание, статус, изображение и дополнительные атрибуты.",
            ),
            (
                "Зачем нужны атрибуты?",
                "Чтобы пользователь мог добавлять собственные характеристики объектов без изменения структуры БД.",
            ),
            (
                "Что такое каталоги?",
                "Это настраиваемые справочники мира: фракции, термины, технологии, социальные группы. У каталогов есть записи, группы, поля и значения.",
            ),
            (
                "Почему PostgreSQL?",
                "Он надежен для реляционных данных, поддерживает индексы, внешние ключи и JSONB, а также хорошо подходит для сложной модели с большим количеством связей.",
            ),
            (
                "Зачем нужны снимки проекта?",
                "Они хранят агрегированное состояние проекта и позволяют быстрее отдавать клиенту или публичной части целостные данные без большого числа отдельных запросов.",
            ),
        ],
    ),
    (
        "Связи и граф",
        [
            (
                "Как работает граф связей?",
                "Объекты становятся узлами графа, а связи из таблиц отношений, владения, персонажных отношений и структур становятся ребрами.",
            ),
            (
                "Где хранятся связи?",
                "В таблицах ObjectRelations, ObjectOwnerships, CharacterRelationships и частично через StructureAssignments.",
            ),
            (
                "Чем связь отличается от графа?",
                "Связь - это запись в БД. Граф - это визуальное представление этих связей на экране.",
            ),
            (
                "Что хранится в ребре графа?",
                "Идентификатор, исходный объект, целевой объект, тип связи, категория связи, а для отношений персонажей также сила, напряженность и признак двусторонности.",
            ),
            (
                "Где хранится расположение узлов графа?",
                "В таблицах RelationGraphLayouts и RelationGraphLayoutItems: там сохраняются координаты, размеры и признак закрепления узлов.",
            ),
            (
                "Как работает таймлайн?",
                "Он хранит события, участников событий, связи между событиями и изменения состояния объектов во времени.",
            ),
        ],
    ),
    (
        "По реализации",
        [
            (
                "Какие технологии использованы?",
                "React, TypeScript, ASP.NET Core Web API, Entity Framework Core, PostgreSQL и Docker Compose.",
            ),
            (
                "Почему клиент-серверная архитектура?",
                "Она разделяет интерфейс и бизнес-логику, упрощает развитие приложения и позволяет централизованно работать с данными.",
            ),
            (
                "Почему React?",
                "Он удобен для динамического интерфейса с карточками, списками, графом связей и таймлайном.",
            ),
            (
                "Почему ASP.NET Core?",
                "Он подходит для надежного API, хорошо работает с PostgreSQL через Entity Framework Core и удобен для серверной бизнес-логики.",
            ),
            (
                "Почему изображения не хранятся прямо в БД?",
                "Так база не раздувается. В БД хранится путь и метаданные файла, а сам файл лежит в файловом хранилище.",
            ),
        ],
    ),
    (
        "По тестированию",
        [
            (
                "Что тестировалось?",
                "Создание проекта, объектов, каталогов, структур, связей, событий таймлайна, загрузка изображений и экспорт DOCX.",
            ),
            (
                "Какие виды тестирования применялись?",
                "Unit-тесты, интеграционные тесты API, тесты клиентских функций и E2E smoke-проверки основных сценариев.",
            ),
            (
                "Как проверялась целостность данных?",
                "Через серверную валидацию, внешние ключи, уникальные ограничения и индексы для основных связей.",
            ),
        ],
    ),
    (
        "Практическая значимость",
        [
            (
                "Кому полезно приложение?",
                "Авторам художественных произведений, сценаристам, разработчикам игровых миров и небольшим творческим группам.",
            ),
            (
                "В чем практическая польза?",
                "Приложение помогает систематизировать материалы проекта, быстрее находить сведения и снижать риск противоречий в художественном мире.",
            ),
            (
                "Что можно развивать дальше?",
                "Совместную работу, расширенный поиск, версионирование, публичные страницы проектов, импорт и экспорт, а также улучшенную визуализацию графов.",
            ),
        ],
    ),
    (
        "Что еще могут спросить",
        [
            (
                "Почему не использовать обычные таблицы или Word?",
                "Потому что они плохо показывают связи и изменения во времени. StoryDB хранит данные структурно и позволяет строить графы и таймлайн.",
            ),
            (
                "Можно ли использовать систему для разных жанров?",
                "Да, потому что атрибуты, каталоги и структуры настраиваются пользователем под конкретный художественный мир.",
            ),
            (
                "Что будет, если удалить проект?",
                "Связанные с ним объекты, каталоги, структуры и события удаляются каскадно согласно настройкам связей.",
            ),
            (
                "Что самое сложное в реализации?",
                "Согласовать гибкую модель данных: объекты, каталоги, структуры, связи и таймлайн так, чтобы они работали вместе и не требовали отдельной схемы под каждый проект.",
            ),
            (
                "Где граница между научной новизной и практической разработкой?",
                "Новизна здесь не в новом классе систем, а в специализированной модели данных и объединении объектов, связей, структур и временных изменений для художественного проекта.",
            ),
            (
                "Как обеспечивается актуальность графа?",
                "Граф строится из текущих данных БД. Если пользователь меняет связи объектов, при следующей загрузке граф собирается уже на основе обновленных связей.",
            ),
            (
                "Что хранится в JSON?",
                "JSON используется точечно, например для снимков проекта и некоторых настраиваемых данных. Основные сущности и связи вынесены в реляционные таблицы.",
            ),
        ],
    ),
]


def set_cell_border(cell, color="D9E2EC", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 14, 8),
        ("Heading 2", 13, "2E74B5", 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def add_question_answer(doc, question, answer):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    q = p.add_run(question)
    q.bold = True
    q.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(7)
    a = p.add_run(answer)
    a.font.size = Pt(10.5)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Шпаргалка: вопросы и ответы по ВКР")
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("2E74B5")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("StoryDB - веб-приложение для авторов художественных произведений")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("555555")

    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, 9936, 120)
    set_cell_width(table.rows[0].cells[0], 2200)
    set_cell_width(table.rows[0].cells[1], 7736)
    for cell in table.rows[0].cells:
        set_cell_border(cell)
    set_cell_shading(table.rows[0].cells[0], "E8EEF5")
    table.rows[0].cells[0].text = "Как пользоваться"
    table.rows[0].cells[1].text = "Учить не дословно, а как опорные формулировки. На защите лучше отвечать коротко и добавлять детали только по уточняющим вопросам."
    for paragraph in table.rows[0].cells[0].paragraphs:
        for run in paragraph.runs:
            run.bold = True

    doc.add_paragraph()

    for section_title, qa_items in SECTIONS:
        doc.add_heading(section_title, level=1)
        for question, answer in qa_items:
            add_question_answer(doc, question, answer)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
