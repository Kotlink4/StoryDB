from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("Речь_к_защите_StoryDB.docx")


SPEECH_PARAGRAPHS = [
    "Уважаемые члены государственной экзаменационной комиссии, вашему вниманию представляется выпускная квалификационная работа на тему: «Разработка веб-приложения для авторов художественных произведений».",
    "Целью работы являлась разработка веб-приложения, которое позволяет авторам художественных произведений вести базу знаний проекта, хранить сведения о художественном мире, анализировать связи между сущностями и фиксировать события во временной линии.",
    "Для достижения цели были поставлены следующие задачи: проанализировать предметную область, сформировать требования к приложению, спроектировать архитектуру и модель данных, реализовать серверную и клиентскую части, а также проверить основные функциональные сценарии работы системы.",
    "В процессе анализа были рассмотрены существующие решения, такие как World Anvil, Campfire, Notebook.ai и Kanka. Эти системы позволяют вести материалы художественного мира, однако часть из них ориентирована на публикацию, часть - на заранее заданные шаблоны или настольные ролевые игры. Разрабатываемое приложение StoryDB ориентировано на более гибкое ведение базы художественного проекта, работу со связями, каталогами, структурами и временной линией.",
    "На контекстной диаграмме показано общее взаимодействие пользователя с системой. Автор работает с веб-приложением через клиентский интерфейс, создает и редактирует проекты, объекты, связи, события и каталоги. Система, в свою очередь, обращается к серверной части, базе данных и файловому хранилищу для сохранения текстовых данных, изображений и результатов экспорта.",
    "Архитектура приложения построена по клиент-серверному принципу. Клиентская часть реализована с использованием React и TypeScript. Серверная часть выполнена на ASP.NET Core Web API. Для хранения данных используется PostgreSQL, а для хранения изображений и других медиафайлов применяется файловое хранилище. Развертывание приложения предусмотрено с использованием Docker Compose.",
    "В основе модели данных находится проект художественного произведения. С ним связаны объекты художественного мира: персонажи, предметы, места, организации и другие сущности. Также модель включает каталоги, структуры, отношения между объектами, события временной линии, медиафайлы и пользовательские раскладки. Такой подход позволяет хранить не только отдельные карточки, но и связи между ними.",
    "Рабочая область приложения построена так, чтобы пользователь мог быстро переходить между основными разделами проекта. Слева расположена навигация по типам данных, в центральной части отображается список объектов, а справа открывается карточка выбранного элемента. Это позволяет сохранять контекст работы и не переходить постоянно между несвязанными страницами.",
    "Редактирование объектов и каталогов выполнено через карточки и формы. Для объекта можно указать основные сведения, характеристики, связи, принадлежность к каталогам и структурам. Каталоги используются для справочников художественного мира, например фракций, терминов, технологий или социальных групп. Это делает модель данных более гибкой и позволяет адаптировать приложение под разные произведения.",
    "Отдельное внимание уделено связям и временной линии. Граф связей позволяет наглядно отображать отношения между объектами: например, родственные связи, членство в организации, владение предметом или территориальную принадлежность. Временная линия используется для фиксации событий, периодов и изменений состояния объектов. Это важно для произведений, где мир меняется по ходу сюжета.",
    "В результате выполнения работы было разработано веб-приложение StoryDB. В нем реализованы база объектов художественного мира, каталоги, структуры, граф связей и таймлайн. Также поддержаны загрузка изображений, снимки проекта, экспорт DOCX-досье, контейнерное развертывание и эксплуатационные эндпоинты. Основные функциональные сценарии были проверены.",
    "Таким образом, цель выпускной квалификационной работы достигнута: разработано веб-приложение для авторов художественных произведений, позволяющее систематизировать материалы проекта, хранить сведения о художественном мире и работать со связями и событиями во времени.",
    "Спасибо за внимание.",
]


def set_cell_border(cell, color="D9E2EC", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, width_dxa, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("Речь к защите ВКР")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_metadata_table(doc):
    table = doc.add_table(rows=3, cols=2)
    set_table_width(table, 9360, 120)
    widths = [2400, 6960]
    rows = [
        ("Тема", "Разработка веб-приложения для авторов художественных произведений"),
        ("Проект", "StoryDB"),
        ("Ориентир", "выступление примерно на 5 минут"),
    ]
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            set_cell_border(cell)
            set_cell_width(cell, widths[col_index])
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
        row.cells[0].text = rows[row_index][0]
        row.cells[1].text = rows[row_index][1]
        set_cell_shading(row.cells[0], "E8EEF5")
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()


def add_slide_cues(doc):
    cues = [
        "1-2. Представиться и назвать тему работы.",
        "3. Сформулировать цель и задачи.",
        "4. Коротко объяснить отличие от аналогов.",
        "5-7. Показать контекст, архитектуру и модель данных.",
        "8-10. Рассказать, как пользователь работает с объектами, каталогами, связями и таймлайном.",
        "11. Подвести результат и завершить выступление.",
    ]
    doc.add_heading("Опорный порядок по слайдам", level=2)
    for cue in cues:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(cue)


def build():
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Речь к защите ВКР")
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("2E74B5")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    sub_run = subtitle.add_run("StoryDB")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor.from_string("555555")

    add_metadata_table(doc)
    add_slide_cues(doc)

    doc.add_heading("Текст выступления", level=2)
    for paragraph in SPEECH_PARAGRAPHS:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(paragraph)
        run.font.name = "Calibri"
        run.font.size = Pt(12)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    note.paragraph_format.space_after = Pt(0)
    note_run = note.add_run("Перед выступлением проверь единое написание ФИО на титульных слайдах презентации.")
    note_run.font.name = "Calibri"
    note_run.font.size = Pt(10)
    note_run.font.italic = True
    note_run.font.color.rgb = RGBColor.from_string("555555")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
