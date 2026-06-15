from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("docs/templates")
OUT_PATH = OUT_DIR / "storydb-dossier-template.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(91, 111, 135)
INK = RGBColor(23, 31, 42)
LIGHT_BLUE = "E8EEF5"
SOFT = "F4F6F9"
BORDER = "CAD6E2"
WHITE = RGBColor(255, 255, 255)

TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_START_END = 120


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=CELL_MARGIN_TOP_BOTTOM, bottom=CELL_MARGIN_TOP_BOTTOM,
                     start=CELL_MARGIN_START_END, end=CELL_MARGIN_START_END):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        border = tc_borders.find(tag)
        if border is None:
            border = OxmlElement(f"w:{edge}")
            tc_borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def add_text(paragraph, text, size=11, color=INK, bold=False, italic=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return run


def add_paragraph(doc, text="", size=11, color=INK, bold=False, italic=False, after=6, before=0,
                  align=None, style=None):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    set_paragraph_spacing(paragraph, before=before, after=after)
    if text:
        add_text(paragraph, text, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(text)
    return paragraph


def set_style(style, size, color, bold=False, before=0, after=6, line=1.25):
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style(styles["Normal"], 11, INK, before=0, after=6, line=1.25)
    set_style(styles["Heading 1"], 16, BLUE, bold=True, before=18, after=10)
    set_style(styles["Heading 2"], 13, BLUE, bold=True, before=14, after=7)
    set_style(styles["Heading 3"], 12, DARK_BLUE, bold=True, before=10, after=5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(header, after=0)
    add_text(header, "StoryDB / шаблон экспорта досье", size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(footer, after=0)
    add_text(footer, "{{ProjectName}} · {{ExportDate}}", size=9, color=MUTED)


def add_cover(doc):
    add_paragraph(doc, "StoryDB", size=12, color=MUTED, bold=True, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        doc,
        "Шаблон экспорта досье",
        size=28,
        color=NAVY,
        bold=True,
        after=4,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        doc,
        "Единый Word-формат для персонажей, организаций, предметов и мест",
        size=13,
        color=MUTED,
        after=26,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    table = doc.add_table(rows=2, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340])
    labels = [
        ("Проект", "{{ProjectName}}"),
        ("Тип экспорта", "{{ExportScope}}"),
        ("Объектов", "{{ObjectCount}}"),
        ("Дата", "{{ExportDate}}"),
        ("Автор", "{{AuthorName}}"),
        ("Версия", "{{TemplateVersion}}"),
        ("Контекст", "{{TimelineContext}}"),
        ("Доступ", "{{Visibility}}"),
    ]
    for cell, (label, value) in zip([cell for row in table.rows for cell in row.cells], labels):
        set_cell_shading(cell, SOFT)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=2)
        add_text(p, label.upper(), size=8.5, color=MUTED, bold=True)
        p2 = cell.add_paragraph()
        set_paragraph_spacing(p2, after=0)
        add_text(p2, value, size=11, color=NAVY, bold=True)

    add_paragraph(
        doc,
        "Назначение шаблона",
        size=12,
        color=DARK_BLUE,
        bold=True,
        before=20,
        after=6,
    )
    add_paragraph(
        doc,
        "Этот документ задаёт визуальный формат для серверного экспорта StoryDB. "
        "Плейсхолдеры в двойных фигурных скобках заменяются данными проекта, объекта и выбранного временного контекста.",
        size=11,
        color=INK,
        after=8,
    )
    add_paragraph(
        doc,
        "Основная идея: каждое досье начинается с компактной паспортной карточки, "
        "после чего идут секции с характеристиками, каталогами, принадлежностью к структурам и связями.",
        size=11,
        color=INK,
        after=8,
    )
    add_paragraph(doc, "Разделы типов", size=12, color=DARK_BLUE, bold=True, before=14, after=6)
    add_paragraph(
        doc,
        "{{ObjectTypeSection}}: Персонажи / Предметы / Места / Организации. "
        "Внутри каждого раздела идут только выбранные объекты этого типа.",
        size=11,
        color=INK,
        after=8,
    )


def add_dossier_page(doc):
    doc.add_section(WD_SECTION_START.NEW_PAGE)

    add_paragraph(doc, "{{ObjectTypeSection}}", size=22, color=NAVY, bold=True, before=8, after=4)
    add_paragraph(doc, "{{ObjectTypeCount}} объектов в этом разделе", size=13, color=MUTED, after=18)

    add_heading(doc, "{{ObjectFullName}}", level=1)
    add_paragraph(doc, "{{ObjectKindLabel}} · {{TimelineContext}}", size=10.5, color=MUTED, after=12)

    head = doc.add_table(rows=1, cols=2)
    set_table_geometry(head, [2500, 6860])
    image_cell = head.cell(0, 0)
    data_cell = head.cell(0, 1)
    set_cell_shading(image_cell, LIGHT_BLUE)
    set_cell_shading(data_cell, "FFFFFF")
    image_p = image_cell.paragraphs[0]
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(image_p, before=48, after=48)
    add_text(image_p, "{{ObjectPortrait}}\nвстроенное изображение объекта", size=11, color=MUTED, bold=True)

    data_p = data_cell.paragraphs[0]
    set_paragraph_spacing(data_p, after=4)
    add_text(data_p, "Досье объекта", size=8.5, color=MUTED, bold=True)
    title = data_cell.add_paragraph()
    set_paragraph_spacing(title, after=6)
    add_text(title, "{{ObjectFullName}}", size=20, color=NAVY, bold=True)
    subtitle = data_cell.add_paragraph()
    set_paragraph_spacing(subtitle, after=8)
    add_text(subtitle, "{{ObjectType}}", size=11, color=MUTED)

    facts = data_cell.add_table(rows=2, cols=3)
    set_table_geometry(facts, [2286, 2287, 2287])
    fact_values = [
        ("Статус", "{{CurrentStatus}}"),
        ("Возраст / время", "{{AgeOrTime}}"),
        ("Роль", "{{Role}}"),
        ("Фамилия / дом", "{{SurnameOrHouse}}"),
        ("Каталог", "{{PrimaryCatalog}}"),
        ("Структура", "{{PrimaryStructure}}"),
    ]
    for cell, (label, value) in zip([cell for row in facts.rows for cell in row.cells], fact_values):
        set_cell_shading(cell, SOFT)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=1)
        add_text(p, label.upper(), size=8, color=MUTED, bold=True)
        p2 = cell.add_paragraph()
        set_paragraph_spacing(p2, after=0)
        add_text(p2, value, size=10.5, color=INK, bold=True)

    add_heading(doc, "Описание", level=2)
    add_paragraph(doc, "{{Description}}", size=11, color=INK)

    add_heading(doc, "Характеристики", level=2)
    attributes = doc.add_table(rows=1, cols=4)
    set_table_geometry(attributes, [1900, 2480, 2480, 2500])
    for cell, text in zip(attributes.rows[0].cells, ["Группа", "Характеристика", "Значение", "Комментарий"]):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0)
        add_text(p, text, size=9, color=NAVY, bold=True)
    for values in [
        ("{{AttributeGroup}}", "{{AttributeName}}", "{{AttributeValue}}", "{{AttributeNote}}"),
        ("{{AttributeGroup}}", "{{AttributeName}}", "{{AttributeValue}}", "{{AttributeNote}}"),
    ]:
        row = attributes.add_row()
        for cell, value in zip(row.cells, values):
            set_cell_margins(cell)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0)
            add_text(p, value, size=9.5, color=INK)

    add_heading(doc, "Каталоги и структуры", level=2)
    catalog = doc.add_table(rows=1, cols=3)
    set_table_geometry(catalog, [2500, 3430, 3430])
    for cell, text in zip(catalog.rows[0].cells, ["Раздел", "Запись", "Контекст"]):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0)
        add_text(p, text, size=9, color=NAVY, bold=True)
    for values in [
        ("{{CatalogName}}", "{{CatalogEntry}}", "{{CatalogContext}}"),
        ("{{StructureName}}", "{{StructureNode}}", "{{StructureRole}}"),
    ]:
        row = catalog.add_row()
        for cell, value in zip(row.cells, values):
            set_cell_margins(cell)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0)
            add_text(p, value, size=9.5, color=INK)

    add_heading(doc, "Связи", level=2)
    add_paragraph(
        doc,
        "{{RelationsSummary}}",
        size=11,
        color=INK,
        after=8,
    )
    relations = doc.add_table(rows=1, cols=4)
    set_table_geometry(relations, [2200, 2200, 2100, 2860])
    for cell, text in zip(relations.rows[0].cells, ["Источник", "Цель", "Тип", "Описание"]):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0)
        add_text(p, text, size=9, color=NAVY, bold=True)
    row = relations.add_row()
    for cell, value in zip(row.cells, ["{{RelationSource}}", "{{RelationTarget}}", "{{RelationType}}", "{{RelationDescription}}"]):
        set_cell_margins(cell)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0)
        add_text(p, value, size=9.5, color=INK)

def add_notes_page(doc):
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "Правила заполнения шаблона", level=1)
    add_paragraph(doc, "Этот лист не обязан попадать в пользовательский экспорт. Он нужен разработчику как карта шаблона.", color=INK)

    notes = doc.add_table(rows=1, cols=3)
    set_table_geometry(notes, [2500, 3430, 3430])
    for cell, text in zip(notes.rows[0].cells, ["Блок", "Плейсхолдеры", "Правило генерации"]):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0)
        add_text(p, text, size=9, color=NAVY, bold=True)
    rows = [
        ("Раздел", "{{ObjectTypeSection}}, {{ObjectTypeCount}}", "Появляется перед первой страницей каждого типа объектов."),
        ("Паспорт", "{{ObjectFullName}}, {{ObjectType}}, {{ObjectPortrait}}", "Всегда в начале каждого досье."),
        ("Контекст", "{{TimelineContext}}, {{CurrentStatus}}", "Подставлять состояние выбранной точки таймлайна."),
        ("Таблицы", "{{AttributesTable}}, {{CatalogsTable}}, {{RelationsTable}}", "Строки добавляются только при наличии данных."),
        ("Пустые секции", "{{EmptyState}}", "Показывать короткую строку 'Данных пока нет', а не оставлять пустой блок."),
    ]
    for values in rows:
        row = notes.add_row()
        for cell, value in zip(row.cells, values):
            set_cell_margins(cell)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0)
            add_text(p, value, size=9.5, color=INK)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    add_dossier_page(doc)
    add_notes_page(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH.resolve())


if __name__ == "__main__":
    build()
